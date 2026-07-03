# -*- coding: utf-8 -*-
"""
习思想题库批量解析脚本
功能：遍历所有章节docx/doc文件，自动识别单选/多选/判断，提取题干、选项、答案
输出：cleaned_questions.json
"""

import os, re, json, sys
from docx import Document

# ---- 项目根目录（脚本在 scripts/ 子目录下） ----
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

# ---- 配置 ----
BANK_DIR = os.path.join(PROJECT_ROOT, "_temp_bank", "习思想题库 - 副本")
OUTPUT_JSON = os.path.join(PROJECT_ROOT, "cleaned_questions.json")

# 章节文件名→章节名映射
CHAPTER_MAP = {
    "0（导论）": "导论",
    "1（第一章）": "第一章 新时代坚持和发展中国特色社会主义",
    "2（第二章）": "第二章 以中国式现代化全面推进中华民族伟大复兴",
    "3.（第三章）": "第三章 坚持党的全面领导",
    "4（第四章）": "第四章 坚持以人民为中心",
    "5（第五章）": "第五章 全面深化改革开放",
    "6.（第六章）": "第六章 推动高质量发展",
    "7（第七章）": "第七章 社会主义现代化建设的教育科技人才战略",
    "8（第八章）": "第八章 发展全过程人民民主",
    "9（第九章）": "第九章 全面依法治国",
    "10（第十章）": "第十章 建设社会主义文化强国",
    "11（第十一章）": "第十一章 以保障和改善民生为重点加强社会建设",
    "12（第十二章）": "第十二章 建设社会主义生态文明",
    "13（第十三章）": "第十三章 维护和塑造国家安全",
    "14（第十四章）": "第十四章 建设巩固国防和强大人民军队",
    "15（第十五章）": '第十五章 坚持"一国两制"和推进祖国完全统一',
    "16（第十六章）": '第十六章 中国特色大国外交和推动构建人类命运共同体',
    "17（第十七章）": "第十七章 全面从严治党",
}


def extract_text_from_docx(filepath):
    """从docx文件提取纯文本行"""
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return lines


def extract_text_from_doc(filepath):
    """
    对于旧版.doc文件，尝试用python-docx读取
    如果失败则用antiword或手动处理
    """
    try:
        return extract_text_from_docx(filepath)
    except Exception:
        # .doc 文件无法直接用 python-docx 解析
        # 尝试用 olefile 或 返回空
        print(f"  [WARN] 无法解析旧版.doc文件: {os.path.basename(filepath)}")
        return []


def classify_section(line):
    """识别题型分节标题，返回题型代码"""
    line_clean = line.replace(" ", "").replace("\u3000", "")
    if re.search(r'一[.．、]\s*单', line_clean) or re.search(r'单[项选]', line_clean):
        return 'single'
    if re.search(r'二[.．、]\s*多', line_clean) or re.search(r'多[项选]', line_clean):
        return 'multiple'
    if re.search(r'三[.．、]\s*判', line_clean) or re.search(r'判[断]', line_clean):
        return 'judge'
    if re.search(r'四[.．、]\s*简', line_clean) or re.search(r'简[答]', line_clean):
        return 'essay'
    return None


def parse_answer_from_title(title):
    """
    从题干中提取嵌入的答案
    格式: （  C   ） 或 （ ABCDE  ） 或 （  √   ）或 （  ×   ）
    返回: (clean_title, answer_str)
    """
    # 匹配全角括号内的答案
    # 模式: （ 单个/多个字母/√/×  ）
    pattern = r'[（(]\s*([A-Z√×✓✗\s]+?)\s*[）)]'
    match = re.search(pattern, title)
    if match:
        answer_raw = match.group(1).strip()
        # 清理题干：移除答案括号
        clean_title = re.sub(r'\s*[（(]\s*[A-Z√×✓✗\s]*\s*[）)]\s*', '', title).strip()
        # 标准化答案
        answer = answer_raw.replace(' ', '').upper()
        if '√' in answer or '✓' in answer:
            answer = '正确'
        elif '×' in answer or '✗' in answer:
            answer = '错误'
        return clean_title, answer
    return title, ''


def parse_options(lines, start_idx):
    """从指定位置开始收集选项行（A. B. C. ... 开头）"""
    options = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        # 匹配 A. 或 A、 或 A． 开头
        if re.match(r'^[A-H][.、．)\s]', line):
            # 去掉选项前缀
            opt_text = re.sub(r'^[A-H][.、．)\s]+', '', line).strip()
            options.append(opt_text)
            i += 1
        else:
            break
    return options, i


def is_question_start(line):
    """判断是否为题目起始行（数字+标点开头）"""
    return bool(re.match(r'^\d+[.、．]', line.strip()))


def parse_chapter(filepath, chapter_name):
    """解析单个章节文件，返回题目列表"""
    print(f"\n📖 解析: {chapter_name}")
    
    # 提取文本
    if filepath.endswith('.docx'):
        lines = extract_text_from_docx(filepath)
    else:
        lines = extract_text_from_doc(filepath)
    
    if not lines:
        print(f"  ⚠️ 无内容，跳过")
        return []
    
    questions = []
    current_type = 'single'  # 默认单选题
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 检查题型分节
        section_type = classify_section(line)
        if section_type:
            current_type = section_type
            i += 1
            continue
        
        # 检查是否为新题目开始
        if not is_question_start(line):
            # 可能是章节标题或其他内容，跳过
            i += 1
            continue
        
        # 开始解析一道题
        # 1. 题干行（含嵌入答案）
        title, answer = parse_answer_from_title(line)
        # 移除题号前缀
        title = re.sub(r'^\d+[.、．]\s*', '', title).strip()
        
        # 2. 收集后续选项
        options, next_idx = parse_options(lines, i + 1)
        
        # 3. 判断题特殊处理
        if current_type == 'judge':
            if not answer:
                # 尝试从题干推断
                title_clean = title.replace(' ', '')
                if '√' in title_clean or '✓' in title_clean:
                    answer = '正确'
                elif '×' in title_clean or '✗' in title_clean:
                    answer = '错误'
            # 判断题选项通常是["正确", "错误"]
            if not options:
                options = ["正确", "错误"]
        
        # 4. 跳过无答案的题（但保留到列表供人工检查）
        if not answer:
            answer = '（待确认）'
        
        # 构建题目对象
        q = {
            "type": current_type,
            "title": title,
            "options": options,
            "answer": answer,
            "analysis": "",
            "chapter": chapter_name,
        }
        questions.append(q)
        
        # 跳到选项结束后的位置
        i = max(next_idx, i + 1)
    
    print(f"  ✅ 解析出 {len(questions)} 道题 (单选:{sum(1 for q in questions if q['type']=='single')}, 多选:{sum(1 for q in questions if q['type']=='multiple')}, 判断:{sum(1 for q in questions if q['type']=='judge')})")
    return questions


def main():
    all_questions = []
    stats = {'single': 0, 'multiple': 0, 'judge': 0, 'essay': 0}
    
    # 遍历所有文件
    files = sorted(os.listdir(BANK_DIR))
    for fname in files:
        if fname.startswith('~$'):
            continue
        fpath = os.path.join(BANK_DIR, fname)
        if not os.path.isfile(fpath):
            continue
        
        # 匹配章节名
        chapter_display = None
        for key, val in CHAPTER_MAP.items():
            if key in fname:
                chapter_display = val
                break
        
        if not chapter_display:
            chapter_display = fname
        
        # 解析
        try:
            qs = parse_chapter(fpath, chapter_display)
            all_questions.extend(qs)
            for q in qs:
                stats[q['type']] = stats.get(q['type'], 0) + 1
        except Exception as e:
            print(f"  ❌ 解析失败: {e}")
    
    # 输出统计
    print(f"\n{'='*50}")
    print(f"📊 解析完成！总计 {len(all_questions)} 道题")
    for t, c in stats.items():
        type_names = {'single': '单选题', 'multiple': '多选题', 'judge': '判断题', 'essay': '简答题'}
        print(f"  {type_names.get(t, t)}: {c} 道")
    
    # 保存JSON
    with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
        json.dump(all_questions, f, ensure_ascii=False, indent=2)
    
    print(f"\n💾 已保存到: {OUTPUT_JSON}")
    print(f"📏 文件大小: {os.path.getsize(OUTPUT_JSON) / 1024:.1f} KB")
    
    # 检查无答案题目
    no_answer = [q for q in all_questions if q['answer'] == '（待确认）']
    if no_answer:
        print(f"\n⚠️ 以下 {len(no_answer)} 道题未提取到答案，请人工检查:")
        for q in no_answer[:5]:
            print(f"  - [{q['chapter']}] {q['title'][:60]}...")


if __name__ == '__main__':
    main()
